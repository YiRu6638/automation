from docx import Document
import os
import openpyxl

# Define the paths
script_dir = os.path.dirname(os.path.abspath(__file__))
data_path = os.path.join(script_dir, "output", "output-record.xlsx")
template_path = os.path.join(script_dir, "template", "EmployeeName_CoA_YA2024.docx")

data = openpyxl.load_workbook(data_path)
data_sheet = data.active

# Define placeholder replacements
placeholders = (
    "{{name}}",
    "{{TIN number}}"
)

# Replace placeholder text while preserving formatting
def replace_placeholder(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)
    
    # Replace placeholders in tables (if needed)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(placeholder, replacement)
                        run.font.name = 'Georgia'

if __name__ == "__main__":
    # Replace placeholders in the document
    for row in range(3, data_sheet.max_row + 1):
        row_data = list(data_sheet.iter_rows(min_row=row, max_row=row, min_col=1, max_col=26, values_only=True))[0]
        
        doc = Document(template_path)
        
        replace_placeholder(doc, "{{name}}", row_data[1])
        replace_placeholder(doc, "{{TIN number}}", row_data[0])

        new_filepath = os.path.join(script_dir, "coa", row_data[1] + "_CoA_YA2024.docx")
        doc.save(new_filepath)
        print(f"Document saved to {new_filepath}")
